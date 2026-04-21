import os
import re
import io
import tempfile
import pdfplumber
import pytesseract
import openpyxl

from pdf2image import convert_from_path
from PIL import Image
from docx import Document as DocxDocument


OCR_LANG = "chi_sim+eng"

def _save_uploaded_to_temp(uploaded_file, suffix: str) -> str:
    uploaded_file.seek(0)
    fd, path = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(uploaded_file.read())
            f.flush()
            os.fsync(f.fileno())
    except Exception:
        # 如果写入失败，清理文件
        try:
            os.remove(path)
        except Exception:
            pass
        raise
    return path


def _count_effective_chars(text: str) -> int:
    """统计有效字符(中英文数字)，用于判断 PDF 是否抽到文本"""
    if not text:
        return 0
    return len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text))


def _parse_txt(uploaded_file) -> str:
    uploaded_file.seek(0)
    return uploaded_file.read().decode("utf-8", errors="ignore")


def _parse_image_ocr(path: str) -> str:
    img = Image.open(path)
    return pytesseract.image_to_string(img, lang=OCR_LANG)


def _parse_docx_bytes(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_vals = []
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    row_vals.append(t)
            if row_vals:
                parts.append(" | ".join(row_vals))

    return "\n".join(parts)


def _parse_xlsx_bytes(data: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"【Sheet】{ws.title}")
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if row_vals:
                parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def _parse_pdf_text_then_ocr(path: str) -> str:
    """
    PDF：优先抽取文本；抽不到(扫描件/图片型PDF)则 OCR
    """
    text_parts = []
    total_effective = 0

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            t = (page.extract_text() or "").strip()
            if t:
                total_effective += _count_effective_chars(t)
                text_parts.append(f"【PDF第{i}页】\n{t}")

    # 阈值：有效字符太少 -> 认为抽不到 -> OCR
    if total_effective >= 80:
        return "\n\n".join(text_parts).strip()

    # OCR fallback（扫描PDF）
    ocr_parts = []
    images = convert_from_path(path, dpi=300)
    for i, img in enumerate(images, start=1):
        t = pytesseract.image_to_string(img, lang=OCR_LANG).strip()
        if t:
            ocr_parts.append(f"【OCR第{i}页】\n{t}")

    return "\n\n".join(ocr_parts).strip()


def parse_file(uploaded_file) -> str:
    """
    支持：txt / png/jpg/jpeg / pdf(抽不到就OCR) / docx / xlsx
    """
    uploaded_file.seek(0)
    filename = (uploaded_file.name or "").lower()

    if filename.endswith(".txt"):
        return _parse_txt(uploaded_file).strip()

    if filename.endswith(".docx"):
        uploaded_file.seek(0)
        data = uploaded_file.read()
        return _parse_docx_bytes(data).strip()

    if filename.endswith(".xlsx"):
        uploaded_file.seek(0)
        data = uploaded_file.read()
        return _parse_xlsx_bytes(data).strip()

    if filename.endswith((".png", ".jpg", ".jpeg")):
        tmp = _save_uploaded_to_temp(uploaded_file, ".png")
        try:
            return _parse_image_ocr(tmp).strip()
        finally:
            os.remove(tmp)

    if filename.endswith(".pdf"):
        tmp = _save_uploaded_to_temp(uploaded_file, ".pdf")
        try:
            return _parse_pdf_text_then_ocr(tmp).strip()
        finally:
            os.remove(tmp)

    return ""
